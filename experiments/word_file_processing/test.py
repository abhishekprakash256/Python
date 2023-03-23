from docx import Document

doc = Document("testFile.docx")

all_paras = doc.paragraphs

for para in all_paras:
    print(para.text)


added = doc.add_paragraph("Microsoft Word (All versions) has a hidden feature that allows you to insert Lorem Ipsum or Random text into your document easily. I actually use the feature all the Word 2007, 2010, and the latest Word 2013 has a feature that allows you to create Random text or Lorem Ipsum text quickly and easily using a simple command. It’s very convenient for bloggers who need to create a mock website with dummy text or create a mockup of something for a PowerPoint presentation.The new para This is the second paragraph of a MS Word file.")


added.add_run("gitlab,ci-cd,redux,ansible")
doc.save("testFile.docx")